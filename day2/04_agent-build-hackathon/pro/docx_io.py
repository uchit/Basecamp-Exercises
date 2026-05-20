"""DOCX import / export.

Customer-grade RFPs arrive as Word documents and have to go back the same
way. This module: (a) parses a .docx into a list of question dicts compatible
with pro_run.py, and (b) writes an answered .docx where each question is
followed by the Pro-generated answer + sources + confidence badge.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


_QUESTION_PATTERN = re.compile(r"^\s*(?:Q[\d\.]+|Question\s+\d+|\d+\.|\d+\))\s*[:\.\-]?\s*", re.IGNORECASE)


@dataclass
class ImportedQuestion:
    id: str
    text: str
    category: str = "general"


# --------------------------------------------------------------------------
# Import
# --------------------------------------------------------------------------

def import_questions(path: str | Path) -> list[ImportedQuestion]:
    """Parse a Word .docx into questions.

    Heuristics:
      - Numbered or "Q#" prefixed paragraphs → new question.
      - Question paragraphs may be tables; each row's first non-empty cell
        treated as the question if it ends with '?' or starts with a marker.
    """
    doc = Document(str(path))
    questions: list[ImportedQuestion] = []
    counter = 0

    # Paragraphs
    current: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            if current:
                joined = " ".join(current).strip()
                if joined.endswith("?") or "?" in joined:
                    counter += 1
                    questions.append(ImportedQuestion(
                        id=f"Q{counter}",
                        text=_QUESTION_PATTERN.sub("", joined).strip(),
                    ))
                current = []
            continue
        if _QUESTION_PATTERN.match(text):
            # Flush previous, start new.
            if current:
                joined = " ".join(current).strip()
                if joined:
                    counter += 1
                    questions.append(ImportedQuestion(
                        id=f"Q{counter}",
                        text=_QUESTION_PATTERN.sub("", joined).strip(),
                    ))
            current = [text]
        else:
            current.append(text)
    if current:
        joined = " ".join(current).strip()
        if joined:
            counter += 1
            questions.append(ImportedQuestion(
                id=f"Q{counter}",
                text=_QUESTION_PATTERN.sub("", joined).strip(),
            ))

    # Tables (RFPs often arrive as Q-in-table format)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if text.endswith("?") and len(text.split()) >= 4:
                    counter += 1
                    questions.append(ImportedQuestion(
                        id=f"Q{counter}",
                        text=text,
                    ))
                    break  # one question per row

    return questions


# --------------------------------------------------------------------------
# Export
# --------------------------------------------------------------------------

_CONF_COLORS = {
    "high":   RGBColor(0x30, 0xB0, 0x4A),
    "medium": RGBColor(0xFF, 0x9F, 0x0A),
    "low":    RGBColor(0xFF, 0x45, 0x3A),
}


def export_responses(report: dict, out_path: str | Path,
                      *, customer_name: str | None = None) -> Path:
    """Write a customer-facing DOCX with question → answer → sources blocks."""
    doc = Document()
    title = doc.add_heading(report.get("rfp_name", "RFP response"), level=0)

    intro = doc.add_paragraph()
    intro.add_run(
        f"This response was prepared by Helios Security."
        + (f" Customer: {customer_name}." if customer_name else "")
    ).italic = True
    doc.add_paragraph()

    for a in report.get("answers", []) or []:
        # Question header
        h = doc.add_heading(level=2)
        h.add_run(f"{a.get('question_id')}: ").bold = True
        h.add_run(a.get("question", ""))

        # Answer body
        doc.add_paragraph(a.get("answer", ""))

        # Sources line
        srcs = a.get("sources") or []
        if srcs:
            s = doc.add_paragraph()
            s.add_run("Sources: ").bold = True
            s.add_run(", ".join(srcs)).italic = True

        # Confidence badge
        conf = (a.get("confidence") or "low").lower()
        c = doc.add_paragraph()
        run = c.add_run(f"[{conf.upper()} CONFIDENCE]")
        run.bold = True
        run.font.color.rgb = _CONF_COLORS.get(conf, RGBColor(0x6E, 0x6E, 0x73))
        run.font.size = Pt(9)

        # Flags (for internal review only — kept in document)
        flags = a.get("flags") or []
        if flags:
            f = doc.add_paragraph()
            f.add_run("Internal flags for review: ").bold = True
            for flag in flags:
                f.add_run(f"• {flag}\n").italic = True

        doc.add_paragraph()  # spacer

    # Provenance block
    prov = report.get("provenance") or {}
    if prov:
        doc.add_heading("Document provenance", level=3)
        p = doc.add_paragraph()
        p.add_run(
            f"Generated by {prov.get('agent_version', 'pro')} "
            f"using {prov.get('model', 'claude-sonnet-4-6')} on {prov.get('ts_utc', '')}\n"
            f"Run ID: {prov.get('run_id', '')}\n"
            f"KB hash: {prov.get('kb_hash', '')}"
        ).font.size = Pt(8)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
