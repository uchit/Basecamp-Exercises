"""Tests for mega-batch J: product surface + feedback loop modules.
All no-API tests. Uses tmp_path for any disk persistence.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

HERE = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(HERE))

from pro import (
    docx_io, diff, feedback, win_loss, tenancy, slack,
    regression, abtest, calibration, audit_db,
)


# ──────────────────────── docx_io ────────────────────────

class TestDocxIO:
    def test_export_then_count_paragraphs(self, tmp_path):
        from docx import Document  # type: ignore
        report = {
            "rfp_name": "Test RFP",
            "answers": [{
                "question_id": "Q1", "question": "What is X?",
                "answer": "X is Y.", "sources": ["Doc A"],
                "confidence": "high", "flags": [],
            }],
            "provenance": {"agent_version": "pro/1.0.0", "model": "x",
                            "run_id": "r", "kb_hash": "abc",
                            "ts_utc": "2026-01-01T00:00:00Z"},
        }
        out = tmp_path / "out.docx"
        docx_io.export_responses(report, out)
        assert out.exists()
        d = Document(str(out))
        text = "\n".join(p.text for p in d.paragraphs)
        assert "Test RFP" in text
        assert "X is Y." in text
        assert "HIGH CONFIDENCE" in text

    def test_import_questions_from_simple_docx(self, tmp_path):
        from docx import Document  # type: ignore
        d = Document()
        d.add_paragraph("Q1: What is your pricing for 500 endpoints?")
        d.add_paragraph("")
        d.add_paragraph("Q2: Do you hold SOC 2 Type II?")
        d.add_paragraph("")
        d.add_paragraph("3. How do you handle EU data residency?")
        path = tmp_path / "in.docx"
        d.save(str(path))
        qs = docx_io.import_questions(path)
        # at least the three above
        texts = [q.text for q in qs]
        assert any("500 endpoints" in t for t in texts)
        assert any("SOC 2" in t for t in texts)
        assert any("EU data" in t for t in texts)


# ──────────────────────── diff ────────────────────────

class TestDiff:
    def test_word_diff_no_change(self):
        pieces = diff.compute_word_diff("same text", "same text")
        assert all(p.op == "equal" for p in pieces)

    def test_word_diff_insertion(self):
        pieces = diff.compute_word_diff("Alpha beta", "Alpha and beta")
        kinds = [p.op for p in pieces]
        assert "insert" in kinds

    def test_word_diff_deletion(self):
        pieces = diff.compute_word_diff("Alpha and beta", "Alpha beta")
        kinds = [p.op for p in pieces]
        assert "delete" in kinds

    def test_render_html_marks_changes(self):
        pieces = diff.compute_word_diff("old answer", "new answer")
        html = diff.render_html(pieces)
        assert "<ins" in html or "<del" in html

    def test_diff_summary_returns_metrics(self):
        pieces = diff.compute_word_diff("a b c", "a x c")
        s = diff.diff_summary(pieces)
        for k in ("chars_inserted", "chars_deleted", "chars_unchanged", "change_ratio"):
            assert k in s


# ──────────────────────── feedback ────────────────────────

class TestFeedback:
    def _seed_run(self, db_path):
        with audit_db.connect(db_path) as conn:
            audit_db.insert_run(conn, run_id="r1", rfp_name="t")

    def test_invalid_verdict_raises(self, tmp_path):
        self._seed_run(tmp_path / "x.db")
        with pytest.raises(ValueError):
            feedback.record(
                feedback.FeedbackEntry(run_id="r1", question_id="Q1",
                                        reviewer="alice", verdict="bogus"),
                db_path=tmp_path / "x.db",
            )

    def test_record_and_list(self, tmp_path):
        db = tmp_path / "x.db"
        self._seed_run(db)
        feedback.record(
            feedback.FeedbackEntry(run_id="r1", question_id="Q1",
                                    reviewer="alice", verdict="approve"),
            db_path=db,
        )
        rows = feedback.list_for_run("r1", db_path=db)
        assert len(rows) == 1
        assert rows[0]["verdict"] == "approve"

    def test_aggregate_counts(self, tmp_path):
        db = tmp_path / "x.db"
        self._seed_run(db)
        for v in ("approve", "approve", "edit", "reject"):
            feedback.record(
                feedback.FeedbackEntry(run_id="r1", question_id="Q1",
                                        reviewer="x", verdict=v, edit_text="e"),
                db_path=db,
            )
        agg = feedback.aggregate(db_path=db)
        assert agg["total"] == 4
        assert agg["by_verdict"]["approve"] == 2


# ──────────────────────── win_loss ────────────────────────

class TestWinLoss:
    def test_record_and_list_outcomes(self, tmp_path):
        db = tmp_path / "x.db"
        with audit_db.connect(db) as conn:
            audit_db.insert_run(conn, run_id="r1", rfp_name="t")
            audit_db.finish_run(conn, run_id="r1", composite_score=95,
                                  total_cost=0.5, total_calls=10)
        win_loss.record_outcome(win_loss.Outcome(run_id="r1", outcome="won"),
                                  db_path=db)
        out = win_loss.list_outcomes(db_path=db)
        assert len(out) == 1
        assert out[0]["outcome"] == "won"

    def test_invalid_outcome_raises(self, tmp_path):
        with pytest.raises(ValueError):
            win_loss.record_outcome(
                win_loss.Outcome(run_id="r1", outcome="garbage"),
                db_path=tmp_path / "x.db",
            )

    def test_win_rate_by_score_band(self, tmp_path):
        db = tmp_path / "x.db"
        seed = [("r1", 95, "won"), ("r2", 92, "won"),
                ("r3", 70, "lost"), ("r4", 50, "lost")]
        with audit_db.connect(db) as conn:
            for rid, score, _ in seed:
                audit_db.insert_run(conn, run_id=rid, rfp_name="t")
                audit_db.finish_run(conn, run_id=rid, composite_score=score,
                                      total_cost=0, total_calls=0)
        for rid, _, oc in seed:
            win_loss.record_outcome(
                win_loss.Outcome(run_id=rid, outcome=oc), db_path=db,
            )
        rates = win_loss.win_rate_by_score_band(db_path=db)
        bands = {r["band"]: r for r in rates}
        # 90-100 band: 2 wins / 2 in band → win_rate = 1.0
        assert bands["90-100"]["win_rate"] == 1.0


# ──────────────────────── tenancy ────────────────────────

class TestTenancy:
    def test_default_returns_bundled_kb(self):
        t = tenancy.load("default")
        assert t.tenant_id == "default"
        assert t.kb_size() > 0

    def test_invalid_id_raises(self):
        with pytest.raises(ValueError):
            tenancy.load("Bad ID With Spaces")

    def test_missing_tenant_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            tenancy.load("never-exists", kbs_dir=tmp_path)

    def test_save_then_load_roundtrip(self, tmp_path):
        kb = {"doc1": {"source": "Test Doc", "content": "Hello world", "tags": ["t"]}}
        tenancy.save_tenant("acme", kb, kbs_dir=tmp_path)
        t = tenancy.load("acme", kbs_dir=tmp_path)
        assert t.tenant_id == "acme"
        assert "doc1" in t.kb

    def test_schema_validation_rejects_missing_source(self, tmp_path):
        bad = tmp_path / "bad.json"
        bad.write_text(json.dumps({"d1": {"content": "no source"}}))
        # save_tenant validates by load, but here use direct load
        with pytest.raises(ValueError):
            tenancy.load("bad", kbs_dir=tmp_path)


# ──────────────────────── slack ────────────────────────

class TestSlack:
    def _mk_report(self, score=97.0, issues=0, blockers=0):
        return {
            "rfp_name": "Test RFP",
            "answers": [{"question_id": f"Q{i}"} for i in range(5)],
            "composite": {"score": score, "source_coverage": 100.0,
                            "confidence_index": 95.0, "grounding_rate": 100.0,
                            "reviewer_clean": 80.0},
            "cost": {"total_cost": 0.12, "total_calls": 16,
                      "total_input_tokens": 25000, "total_output_tokens": 4000,
                      "wall_clock_s": 30.0, "by_stage": {}},
            "review": {"issues": [{"severity": "blocker"}] * blockers
                       + [{"severity": "warning"}] * (issues - blockers)},
        }

    def test_build_payload_includes_blocks(self):
        p = slack.build_payload(self._mk_report())
        assert "blocks" in p
        assert any(b["type"] == "header" for b in p["blocks"])

    def test_payload_calls_out_blockers(self):
        p = slack.build_payload(self._mk_report(score=60.0, issues=2, blockers=1))
        text = json.dumps(p)
        assert "Blockers to resolve" in text

    def test_post_returns_dict_with_status_when_no_url(self, monkeypatch):
        monkeypatch.delenv("SLACK_WEBHOOK_URL", raising=False)
        r = slack.post(slack.build_payload(self._mk_report()))
        assert r["posted"] is False


# ──────────────────────── regression ────────────────────────

class TestRegression:
    def _mk_report(self, qid, conf, sources, grounded, flags):
        return {
            "answers": [{
                "question_id": qid, "confidence": conf,
                "sources": sources,
                "verification": {"fully_grounded": grounded},
                "flags": flags,
            }],
        }

    def test_update_then_clean_check_returns_no_regressions(self, tmp_path):
        snap_path = tmp_path / "snap.json"
        report = self._mk_report("Q1", "high", ["A"], True, [])
        regression.update_snapshot(report, rfp_key="r1", path=snap_path)
        assert regression.check(report, rfp_key="r1", path=snap_path) == []

    def test_confidence_drop_detected(self, tmp_path):
        snap = tmp_path / "snap.json"
        regression.update_snapshot(
            self._mk_report("Q1", "high", ["A"], True, []),
            rfp_key="r1", path=snap,
        )
        regressed = self._mk_report("Q1", "low", ["A"], True, [])
        regs = regression.check(regressed, rfp_key="r1", path=snap)
        assert any(r.kind == "confidence_dropped" for r in regs)

    def test_lost_source_detected(self, tmp_path):
        snap = tmp_path / "snap.json"
        regression.update_snapshot(
            self._mk_report("Q1", "high", ["A", "B"], True, []),
            rfp_key="r1", path=snap,
        )
        regs = regression.check(
            self._mk_report("Q1", "high", ["A"], True, []),
            rfp_key="r1", path=snap,
        )
        assert any(r.kind == "lost_source" for r in regs)

    def test_grounding_loss_detected(self, tmp_path):
        snap = tmp_path / "snap.json"
        regression.update_snapshot(
            self._mk_report("Q1", "high", ["A"], True, []),
            rfp_key="r1", path=snap,
        )
        regs = regression.check(
            self._mk_report("Q1", "high", ["A"], False, []),
            rfp_key="r1", path=snap,
        )
        assert any(r.kind == "grounding_dropped" for r in regs)


# ──────────────────────── abtest ────────────────────────

class TestABTest:
    def test_control_variant_exists(self):
        v = abtest.assign(forced="control")
        assert v.name == "control"

    def test_forced_unknown_raises(self):
        with pytest.raises(ValueError):
            abtest.assign(forced="not-a-variant")

    def test_deterministic_by_run_id(self):
        v1 = abtest.assign(run_id="run-abc")
        v2 = abtest.assign(run_id="run-abc")
        assert v1.name == v2.name

    def test_disable_drops_from_active(self):
        # register an experimental variant + disable it
        abtest.register(abtest.Variant(name="experimental", model="x", weight=1.0))
        assert any(v.name == "experimental" for v in abtest.list_active_variants())
        abtest.disable("experimental")
        assert not any(v.name == "experimental" for v in abtest.list_active_variants())


# ──────────────────────── calibration ────────────────────────

class TestCalibration:
    def test_cold_start_uses_prior(self, tmp_path):
        db = tmp_path / "x.db"
        r = calibration.calibrate("high", db_path=db)
        # Empty DB → bootstrapped, calibrated ≈ prior (0.90)
        assert r.bucket.is_bootstrapped
        assert abs(r.bucket.calibrated - 0.90) < 0.01

    def test_calibrate_all_buckets(self, tmp_path):
        buckets = calibration.calibrate_all(db_path=tmp_path / "x.db")
        assert set(buckets) == {"high", "medium", "low"}
        for b in buckets.values():
            assert 0 <= b.calibrated <= 1

    def test_advice_string_present(self, tmp_path):
        r = calibration.calibrate("low", db_path=tmp_path / "x.db")
        assert isinstance(r.advice, str)
        assert len(r.advice) > 10

    def test_high_history_updates_calibrated_toward_raw(self, tmp_path):
        db = tmp_path / "x.db"
        with audit_db.connect(db) as conn:
            audit_db.insert_run(conn, run_id="r1", rfp_name="t")
            for i in range(30):
                audit_db.insert_answer(conn, run_id="r1", answer={
                    "question_id": f"Q{i}", "confidence": "high",
                    "sources": ["A"], "answer": "...", "flags": [],
                    "verification": {"fully_grounded": True},
                })
                audit_db.insert_feedback(conn, run_id="r1", question_id=f"Q{i}",
                                          reviewer="r", verdict="approve")
        r = calibration.calibrate("high", db_path=db)
        # Many samples + all approved → calibrated should approach 1.0
        assert r.bucket.calibrated > 0.95
        assert not r.bucket.is_bootstrapped
